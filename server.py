"""
Meeting Automation MCP Server

Claude 오케스트레이션 방식으로 Fireflies + Asana + Notion 워크플로우 가이드 제공
각 MCP 도구를 순차적으로 호출하도록 Claude에게 안내합니다.
"""

from typing import Any
import asyncio
from mcp.server import Server
from mcp.types import Tool, TextContent
import mcp.server.stdio

# 서버 인스턴스
app = Server("meeting-automation")

# 도구 정의
@app.list_tools()
async def list_tools() -> list[Tool]:
    """워크플로우 가이드 도구 목록"""
    return [
        Tool(
            name="guide_meeting_search",
            description="""
회의 검색 가이드

필요한 MCP 도구:
1. Fireflies:search - 날짜나 키워드로 회의 검색

사용 예시:
- 날짜 검색: Fireflies:search query="from:2024-10-24 to:2024-10-24"
- 키워드 검색: Fireflies:search query='keyword:"프로젝트"'
- 최근 회의: Fireflies:search query="limit:10"

검색 결과에서 meeting ID를 확인하세요.
            """,
            inputSchema={
                "type": "object",
                "properties": {
                    "query_hint": {
                        "type": "string",
                        "description": "검색 조건 힌트 (날짜, 키워드 등)"
                    }
                }
            }
        ),
        Tool(
            name="guide_meeting_details",
            description="""
회의 상세정보 조회 가이드

필요한 MCP 도구:
1. Fireflies:get_summary - 회의 요약, 액션아이템, 키워드 가져오기
2. Fireflies:get_transcript - 전체 대화 내용 가져오기 (선택사항)

사용 순서:
1. meeting_id 준비 (검색으로 얻은 ID)
2. Fireflies:get_summary transcriptId="{meeting_id}"
3. (필요시) Fireflies:get_transcript transcriptId="{meeting_id}"

결과에서:
- summary.action_items: 액션아이템 목록
- summary.overview: 회의 요약
- summary.keywords: 주요 키워드
            """,
            inputSchema={
                "type": "object",
                "properties": {
                    "meeting_id": {
                        "type": "string",
                        "description": "Fireflies 회의 ID"
                    }
                },
                "required": ["meeting_id"]
            }
        ),
        Tool(
            name="guide_create_asana_tasks",
            description="""
Asana 태스크 생성 가이드

필요한 MCP 도구:
1. Asana:asana_typeahead_search - 프로젝트 검색
2. Asana:asana_create_task - 태스크 생성

사용 순서:
1. 프로젝트 찾기:
   Asana:asana_typeahead_search 
   resource_type="project"
   query="프로젝트명"
   workspace_gid="{workspace_id}"

2. 액션아이템마다 태스크 생성:
   Asana:asana_create_task
   name="액션아이템 내용"
   project_id="{project_gid}"
   notes="담당자: XXX\n회의: {회의제목}"
   due_on="2024-11-01"

팁:
- 액션아이템에서 담당자 파싱
- 우선순위 설정 가능
- 섹션 지정 가능
            """,
            inputSchema={
                "type": "object",
                "properties": {
                    "action_items": {
                        "type": "string",
                        "description": "Fireflies에서 가져온 액션아이템 텍스트"
                    },
                    "project_name": {
                        "type": "string",
                        "description": "Asana 프로젝트명"
                    }
                }
            }
        ),
        Tool(
            name="guide_save_to_notion",
            description="""
Notion 회의록 저장 가이드

필요한 MCP 도구:
1. Notion:notion-search - 데이터베이스 찾기
2. Notion:notion-create-pages - 페이지 생성

사용 순서:
1. 데이터베이스 찾기:
   Notion:notion-search
   query="회의록" (또는 데이터베이스명)
   query_type="internal"

2. 페이지 생성:
   Notion:notion-create-pages
   parent={data_source_id: "{db_id}"}
   pages=[{
     properties: {
       "제목": "회의제목 - 날짜",
       "날짜": "2024-10-24",
       "참석자": "...",
     },
     content: "# 회의 요약\n...\n# 액션아이템\n..."
   }]

팁:
- 데이터베이스 스키마 먼저 확인 (fetch로)
- 날짜 속성은 date: 형식 사용
- Markdown으로 내용 작성
            """,
            inputSchema={
                "type": "object",
                "properties": {
                    "meeting_summary": {
                        "type": "string",
                        "description": "회의 요약 내용"
                    },
                    "database_name": {
                        "type": "string",
                        "description": "Notion 데이터베이스명"
                    }
                }
            }
        ),
        Tool(
            name="guide_full_workflow",
            description="""
전체 워크플로우 실행 가이드

완전 자동화 순서:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

단계 1: 회의 검색
→ Fireflies:search query="from:YYYY-MM-DD to:YYYY-MM-DD"
→ meeting_id 확인

단계 2: 회의 내용 가져오기  
→ Fireflies:get_summary transcriptId="{meeting_id}"
→ action_items, overview, keywords 추출

단계 3: Asana 태스크 생성
→ Asana:asana_typeahead_search (프로젝트 찾기)
→ 액션아이템 파싱 (담당자, 기한 추출)
→ Asana:asana_create_task (각 아이템마다)

단계 4: Notion 문서화
→ Notion:notion-search (데이터베이스 찾기)
→ Notion:notion-create-pages (회의록 작성)

단계 5: 결과 보고
→ 생성된 태스크 수
→ Notion 페이지 URL
→ 액션아이템 담당자별 정리

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

주의사항:
- 각 단계마다 결과 확인
- 오류 발생시 해당 단계만 재시도
- workspace_gid, database_id 미리 확인 필요
            """,
            inputSchema={
                "type": "object",
                "properties": {
                    "meeting_date": {
                        "type": "string",
                        "description": "회의 날짜 (YYYY-MM-DD)"
                    },
                    "asana_project": {
                        "type": "string",
                        "description": "Asana 프로젝트명"
                    },
                    "notion_database": {
                        "type": "string",
                        "description": "Notion 데이터베이스명"
                    }
                }
            }
        ),
        Tool(
            name="guide_quick_commands",
            description="""
빠른 실행 명령어 모음

📋 회의 검색:
"10/24 회의 찾아줘"
→ Fireflies:search query="from:2024-10-24 to:2024-10-24"

📊 액션아이템 확인:
"이 회의 액션아이템 보여줘" 
→ Fireflies:get_summary transcriptId="{id}"

✅ 태스크 생성:
"액션아이템을 {프로젝트}에 태스크로 만들어"
→ Asana:asana_typeahead_search + asana_create_task

📝 문서화:
"{데이터베이스}에 회의록 저장해줘"
→ Notion:notion-search + notion-create-pages

🚀 전체 자동화:
"10/24 회의 완전 자동화해줘"
→ guide_full_workflow 참조

💡 팁:
- workspace_gid: Asana:asana_list_workspaces
- database_id: Notion:notion-search
- 담당자 ID: Asana:asana_typeahead_search resource_type="user"
            """,
            inputSchema={
                "type": "object",
                "properties": {}
            }
        ),
        Tool(
            name="guide_create_word_document",
            description="""
Word 문서 생성 가이드 (OK금융그룹 양식)

필요한 MCP 도구:
1. desktop-commander - Python 스크립트 실행

사용 순서:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1. 템플릿 파일 준비
   - OK금융그룹 회의록 양식 .docx 파일
   - 경로: /mnt/user-data/uploads/OK금융그룹_회의록__1_.docx

2. Python 스크립트 작성 (python-docx 사용)
   bash_tool로 실행:
   
   python3 << 'PYTHON_EOF'
   from docx import Document
   
   # 템플릿 로드
   doc = Document('/mnt/user-data/uploads/OK금융그룹_회의록__1_.docx')
   
   # 작성일/작성자 수정
   doc.paragraphs[2].text = '작성일 : YYYY. MM. DD  /  작성자 : 이름'
   
   # 첫 번째 테이블 (회의 정보)
   table1 = doc.tables[0]
   table1.rows[0].cells[1].text = '회의일시'
   table1.rows[0].cells[3].text = '회의장소'
   table1.rows[1].cells[1].text = '주관부서'
   table1.rows[1].cells[3].text = '참석부서'
   table1.rows[2].cells[1].text = '참석자'
   table1.rows[3].cells[1].text = '회의안건'
   
   # 두 번째 테이블 (회의내용/결정사항)
   table2 = doc.tables[1]
   table2.rows[1].cells[0].text = '회의 내용'
   table2.rows[3].cells[0].text = '결정사항'
   
   # 세 번째 테이블 (기타)
   table3 = doc.tables[2]
   table3.rows[1].cells[0].text = '기타사항'
   
   # 저장
   doc.save('/mnt/user-data/outputs/회의록_{날짜}_{제목}.docx')
   print('완료')
   PYTHON_EOF

3. 문서 구조
   - 테이블 0: 회의 기본정보 (4행 4열)
     * Row 0: 회의일시, 회의장소
     * Row 1: 주관부서, 참석부서
     * Row 2: 참석자
     * Row 3: 회의안건
   
   - 테이블 1: 회의내용/결정사항 (4행 1열)
     * Row 1: 회의내용
     * Row 3: 결정사항
   
   - 테이블 2: 기타 (2행 1열)
     * Row 1: 기타사항

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

주의사항:
- bash_tool 사용 (Windows desktop-commander와 경로 호환성)
- /mnt/user-data/uploads: 업로드 파일
- /mnt/user-data/outputs: 생성 파일
- python-docx 라이브러리 필요

예제: Fireflies 회의 → Word 문서
1. Fireflies:get_summary로 회의 정보 가져오기
2. 위 스크립트로 양식에 내용 채우기
3. /mnt/user-data/outputs에 저장
            """,
            inputSchema={
                "type": "object",
                "properties": {
                    "meeting_data": {
                        "type": "string",
                        "description": "회의 정보 (제목, 날짜, 참석자, 내용 등)"
                    },
                    "template_path": {
                        "type": "string",
                        "description": "템플릿 파일 경로",
                        "default": "/mnt/user-data/uploads/OK금융그룹_회의록__1_.docx"
                    }
                }
            }
        )
    ]

# 도구 실행
@app.call_tool()
async def call_tool(name: str, arguments: Any) -> list[TextContent]:
    """가이드 도구 실행 핸들러"""
    
    if name == "guide_meeting_search":
        query_hint = arguments.get("query_hint", "")
        
        guide = f"""
🔍 회의 검색 가이드

검색 조건: {query_hint}

다음 도구를 사용하세요:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Fireflies:search
query="{query_hint}"

또는 구체적인 검색어:
- 날짜: "from:2024-10-24 to:2024-10-24"
- 키워드: 'keyword:"프로젝트회의"'
- 최근: "limit:20"

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

검색 후 meeting_id를 확인하고 다음 단계로 진행하세요.
        """
        
        return [TextContent(type="text", text=guide)]
    
    elif name == "guide_meeting_details":
        meeting_id = arguments.get("meeting_id", "")
        
        guide = f"""
📊 회의 상세정보 조회 가이드

Meeting ID: {meeting_id}

1️⃣ 요약 정보 가져오기:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Fireflies:get_summary
transcriptId="{meeting_id}"

결과에서 확인:
✓ summary.action_items - 액션아이템
✓ summary.overview - 회의 요약
✓ summary.keywords - 주요 키워드
✓ title - 회의 제목
✓ dateString - 회의 날짜

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

2️⃣ (선택) 전체 대화 내용:

Fireflies:get_transcript
transcriptId="{meeting_id}"
        """
        
        return [TextContent(type="text", text=guide)]
    
    elif name == "guide_create_asana_tasks":
        action_items = arguments.get("action_items", "")
        project_name = arguments.get("project_name", "")
        
        guide = f"""
✅ Asana 태스크 생성 가이드

프로젝트: {project_name}
액션아이템: {action_items[:200]}...

실행 순서:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1️⃣ Workspace ID 확인:
Asana:asana_list_workspaces

2️⃣ 프로젝트 찾기:
Asana:asana_typeahead_search
resource_type="project"
query="{project_name}"
workspace_gid="{{workspace_id}}"

3️⃣ 각 액션아이템마다 태스크 생성:
Asana:asana_create_task
name="액션아이템 내용"
project_id="{{project_gid}}"
notes="담당자: XXX\\n회의일: YYYY-MM-DD"
due_on="YYYY-MM-DD"

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

💡 팁:
- 액션아이템에서 "담당자:" 패턴 찾기
- 날짜 정보 파싱
- 우선순위 추가 가능
        """
        
        return [TextContent(type="text", text=guide)]
    
    elif name == "guide_save_to_notion":
        meeting_summary = arguments.get("meeting_summary", "")
        database_name = arguments.get("database_name", "")
        
        guide = f"""
📝 Notion 회의록 저장 가이드

데이터베이스: {database_name}
회의 요약: {meeting_summary[:200]}...

실행 순서:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1️⃣ 데이터베이스 찾기:
Notion:notion-search
query="{database_name}"
query_type="internal"

2️⃣ (중요) 데이터베이스 구조 확인:
Notion:notion-fetch
id="{{database_url}}"

→ 속성명 확인 (제목, 날짜, 참석자 등)

3️⃣ 페이지 생성:
Notion:notion-create-pages
parent={{
  data_source_id: "{{collection_id}}"
}}
pages=[{{
  properties: {{
    "제목": "회의명 - 날짜",
    "date:날짜:start": "2024-10-24",
    "참석자": "..."
  }},
  content: "# 회의 요약\\n...\\n\\n# 액션아이템\\n..."
}}]

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

💡 주의:
- data_source_id는 collection:// URL에서 추출
- 날짜 속성은 date:{속성명}:start 형식
- content는 Markdown 형식
        """
        
        return [TextContent(type="text", text=guide)]
    
    elif name == "guide_full_workflow":
        meeting_date = arguments.get("meeting_date", "")
        asana_project = arguments.get("asana_project", "")
        notion_database = arguments.get("notion_database", "")
        
        guide = f"""
🚀 전체 워크플로우 실행 가이드

설정:
- 회의 날짜: {meeting_date}
- Asana 프로젝트: {asana_project}
- Notion 데이터베이스: {notion_database}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

📋 STEP 1: 회의 검색
Fireflies:search
query="from:{meeting_date} to:{meeting_date}"

→ meeting_id 저장

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

📊 STEP 2: 회의 내용 가져오기
Fireflies:get_summary
transcriptId="{{meeting_id}}"

→ action_items, overview, keywords 저장

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

✅ STEP 3: Asana 태스크 생성

3-1. Workspace 확인:
Asana:asana_list_workspaces

3-2. 프로젝트 찾기:
Asana:asana_typeahead_search
resource_type="project"
query="{asana_project}"
workspace_gid="{{workspace_id}}"

3-3. 액션아이템 파싱 및 태스크 생성:
각 액션아이템마다:
Asana:asana_create_task
name="[액션아이템 내용]"
project_id="{{project_gid}}"
notes="담당자: [추출된 담당자]\\n회의: {{회의제목}}"
due_on="[기한]"

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

📝 STEP 4: Notion 문서화

4-1. 데이터베이스 찾기:
Notion:notion-search
query="{notion_database}"
query_type="internal"

4-2. 스키마 확인:
Notion:notion-fetch
id="{{database_url}}"

4-3. 페이지 생성:
Notion:notion-create-pages
parent={{data_source_id: "{{collection_id}}"}}
pages=[{{
  properties: {{
    "[제목속성명]": "{{회의제목}} - {meeting_date}",
    "date:[날짜속성]:start": "{meeting_date}"
  }},
  content: "
# 회의 요약
{{overview}}

# 액션아이템
{{action_items}}

# 주요 키워드
{{keywords}}
"
}}]

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

📈 STEP 5: 결과 정리

생성 완료:
✓ Asana 태스크 {{count}}개
✓ Notion 페이지 1개

담당자별 액션아이템:
[액션아이템을 담당자별로 그룹핑하여 표시]

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

각 단계를 순서대로 실행하세요!
        """
        
        return [TextContent(type="text", text=guide)]
    
    elif name == "guide_quick_commands":
        guide = """
⚡ 빠른 실행 명령어 모음

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

🔍 자주 쓰는 검색:

"10/24 회의 찾아줘"
→ Fireflies:search query="from:2024-10-24 to:2024-10-24"

"최근 회의 10개 보여줘"
→ Fireflies:search query="limit:10"

"프로젝트 회의 검색"
→ Fireflies:search query='keyword:"프로젝트"'

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

📊 정보 조회:

"이 회의 요약해줘"
→ Fireflies:get_summary transcriptId="{id}"

"액션아이템 보여줘"
→ 위 결과에서 summary.action_items

"전체 대화 내용 보여줘"
→ Fireflies:get_transcript transcriptId="{id}"

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

✅ 태스크 생성:

"액션아이템을 Asana 태스크로 만들어"
1. Asana:asana_list_workspaces
2. Asana:asana_typeahead_search (프로젝트)
3. Asana:asana_create_task (각 항목)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

📝 문서화:

"회의록 데이터베이스에 저장해줘"
1. Notion:notion-search query="회의록"
2. Notion:notion-fetch (스키마 확인)
3. Notion:notion-create-pages

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

🚀 완전 자동화:

"10/24 회의 완전 자동화해줘"
→ guide_full_workflow 도구 사용
→ 모든 단계 순차 실행

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

💡 필수 ID 찾기:

Workspace ID:
→ Asana:asana_list_workspaces

Database ID:
→ Notion:notion-search query="데이터베이스명"

담당자 ID:
→ Asana:asana_typeahead_search 
   resource_type="user"
   query="이름"
        """
        
        return [TextContent(type="text", text=guide)]
    
    else:
        raise ValueError(f"알 수 없는 도구: {name}")

# 서버 실행
async def main():
    """서버 시작"""
    async with mcp.server.stdio.stdio_server() as (read_stream, write_stream):
        await app.run(
            read_stream,
            write_stream,
            app.create_initialization_options()
        )

if __name__ == "__main__":
    asyncio.run(main())
